import requests
import streamlit as st
from fpdf import FPDF
from docx import Document

OPENROUTER_KEY = st.secrets["OPENROUTER_API_KEY"]

def gerar_post_blog(pergunta, resposta):
    prompt = f"""
Com base na seguinte resposta técnica gerada a partir de materiais aprendidos:

\"\"\"{resposta}\"\"\"

E considerando a pergunta original:
\"{pergunta}\"

Gere um conteúdo para blog seguindo este formato:
- Um título criativo e objetivo
- Uma introdução curta que gere empatia e conexão
- Subtópicos claros (use H2 e H3)
- Conclusão com um CTA (chamada para ação)
- Linguagem humana, leve, mas profissional
- Se possível, adicione emojis e exemplos reais

Resposta:
"""

    headers = {
        "Authorization": f"Bearer {OPENROUTER_KEY}",
        "Content-Type": "application/json"
    }

    body = {
        "model": "gpt-3.5-turbo",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.8
    }

    response = requests.post("https://openrouter.ai/api/v1/chat/completions", json=body, headers=headers)

    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"]
    else:
        st.error(f"Erro ao gerar post: {response.status_code}")
        st.text(response.text)
        return "Erro ao gerar conteúdo."

# Função para gerar arquivo DOCX
def gerar_docx(post):
    doc = Document()
    doc.add_heading('Conteúdo do Blog', 0)
    doc.add_paragraph(post)
    
    # Salvar o arquivo no caminho adequado
    file_path = "/tmp/conteudo_blog.docx"  # Usando um caminho temporário válido no ambiente
    doc.save(file_path)
    
    return file_path

# Função para gerar arquivo PDF
def gerar_pdf(post):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Certificar-se de que o conteúdo está sendo tratado como UTF-8
    post = post.encode('latin-1', 'replace').decode('latin-1')  # Tentar substituir caracteres incompatíveis com latin1

    pdf.multi_cell(0, 10, post)
    file_path = "/tmp/conteudo_blog.pdf"  # Garantir que o caminho seja válido no ambiente de execução
    pdf.output(file_path)

    return file_path